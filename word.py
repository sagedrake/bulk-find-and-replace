from docx import Document
import eventlog


def docx_find_and_replace(filepath, old_word, new_word):
    """
    Replace instances of the given old word with the given new word in given Word document
    :param filepath: The name of the file to be edited
    :param old_word: The word to be searched for, assumed to be capitalized (e.g. 'Shark' not 'shark')
    :param new_word: The word to replace old_word with, assumed to be capitalized (e.g. 'Whale' not 'WHALE')
    :return: None
    """
    document = Document(filepath)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if old_word.lower() in run.text.lower():
                run.text = str_find_and_replace(run.text, old_word, new_word)

    for s in document.sections:
        subsection_find_and_replace(s.header, old_word, new_word)
        subsection_find_and_replace(s.footer, old_word, new_word)
        subsection_find_and_replace(s.first_page_header, old_word, new_word)
        subsection_find_and_replace(s.first_page_footer, old_word, new_word)
        subsection_find_and_replace(s.even_page_header, old_word, new_word)
        subsection_find_and_replace(s.even_page_footer, old_word, new_word)

    for table in document.tables:
        for col in table.columns:
            for cell in col.cells:
                subsection_find_and_replace(cell, old_word, new_word)

    document.save(filepath)
    eventlog.log_event("Replaced instances of " + old_word + " within" + filepath)


def subsection_find_and_replace(subsection, old_word, new_word):
    for paragraph in subsection.paragraphs:
        for run in paragraph.runs:
            if old_word.lower() in run.text.lower():
                run.text = str_find_and_replace(run.text, old_word, new_word)


def str_find_and_replace(string, old_word, new_word):
    """
    Replace instances of the given old word with the given new word in the given string
    :param string: The string to be searched within
    :param old_word: The word to search for
    :param new_word: The word to replace instances of the old word with
    :return: The inputted string with instances of the old word replaced with instances of the new word
    """
    # old -> new
    lowercase_replaced = string.replace(old_word.lower(), new_word.lower())
    # OLD -> NEW
    uppercase_replaced = lowercase_replaced.replace(old_word.upper(), new_word.upper())
    # Old -> New
    capitalized_replaced = uppercase_replaced.replace(old_word, new_word)

    return capitalized_replaced
